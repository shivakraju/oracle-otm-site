#!/usr/bin/env python3
"""Create a Word template for Oracle OTM blog posts."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_heading(doc, text, level=1, color=RGBColor(0x1e, 0x29, 0x3b)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_instruction(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.italic = True
    return p

# ── Title ──────────────────────────────────────────────
add_heading(doc, "Learn Oracle OTM — New Post Template", level=1)
add_instruction(doc, "Fill in the sections below and send this document. The post will be created for you.")

doc.add_paragraph()

# ── Section 1: Post Details ────────────────────────────
add_heading(doc, "1.  Post Details", level=2)
add_instruction(doc, "Fill in the fields below. These control how the post appears on the website.")

tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
cells = [
    ("Post Title", "e.g.  OTM Rate Management Overview"),
    ("Sidebar Position\n(where should it appear in the topics list?)",
     "e.g.  After 'Basic OTM Configurations'  OR  At the bottom"),
    ("Tags (optional, comma separated)", "e.g.  Rates, Carrier, Configuration"),
]
for i, (label, hint) in enumerate(cells):
    tbl.rows[i].cells[0].text = label
    tbl.rows[i].cells[1].text = hint
    tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# ── Section 2: Post Content ────────────────────────────
add_heading(doc, "2.  Post Content", level=2)
add_instruction(doc, "Write your content below. Follow the formatting rules in Section 3.")

doc.add_paragraph()

# Sample intro paragraph
add_para(doc, "Introduction paragraph — Write a brief overview of the topic here.", size=11)
doc.add_paragraph()

# Sample term definition
p = doc.add_paragraph()
r1 = p.add_run("Term Name: ")
r1.bold = True
r1.font.size = Pt(11)
r2 = p.add_run("Definition or explanation goes here.")
r2.font.size = Pt(11)

doc.add_paragraph()

# Sample step
p = doc.add_paragraph()
r1 = p.add_run("Step 1: ")
r1.bold = True
r1.font.size = Pt(11)
r2 = p.add_run("Description of the step.")
r2.font.size = Pt(11)

doc.add_paragraph()

# Sample bullet list
add_para(doc, "Bullet list example:", bold=True, size=11)
for item in ["First item", "Second item", "Third item"]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Sample code block instruction
add_para(doc, "Code / SQL block example:", bold=True, size=11)
add_instruction(doc, "For code or SQL — use Courier New font in your Word doc so it is easy to identify.")
p = doc.add_paragraph()
run = p.add_run("SELECT * FROM location WHERE domain_name = 'WHD';")
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

doc.add_paragraph()

# Sample image instruction
add_para(doc, "Images:", bold=True, size=11)
add_instruction(doc,
    "Insert images directly into the Word document where you want them to appear. "
    "They will be extracted and placed in the post automatically.")

doc.add_paragraph()

# ── Section 3: Formatting Rules ───────────────────────
add_heading(doc, "3.  Formatting Rules (Quick Reference)", level=2)

rules = [
    ("Term definitions",   "Write  'TermName: explanation'  — the word before : will be made bold automatically."),
    ("Steps",              "Write  'Step 1: description'  — same rule applies."),
    ("Bullet points",      "Use Word's bullet list (or start lines with  *  or  -)."),
    ("Headings",           "Use Word Heading 2 or Heading 3 styles for sub-sections."),
    ("Bold text",          "Use Word bold  (Ctrl+B)  as normal."),
    ("Code / SQL / XML",   "Use Courier New font for any code — it will be placed in a code block on the site."),
    ("Images",             "Insert images inline in the document where you want them."),
    ("Notes / Tips",       "Start a paragraph with  'Note:'  or  'Tip:'  and it will be highlighted on the site."),
]

tbl2 = doc.add_table(rows=len(rules), cols=2)
tbl2.style = 'Table Grid'
for i, (label, rule) in enumerate(rules):
    tbl2.rows[i].cells[0].text = label
    tbl2.rows[i].cells[1].text = rule
    tbl2.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
add_instruction(doc, "That's it! Fill in sections 1 and 2, save the file, and send it.")

# Save
out = "OTM_Post_Template.docx"
doc.save(out)
print(f"Template saved: {out}")
